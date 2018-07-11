import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from textblob import TextBlob
import xlwt

def find_component(t):
    regx1 = re.compile(r"A Requester")
    regx10 = re.compile(r"Requester shall")
    regx11 = re.compile(r"All Requesters shall")
    regx12 = re.compile(r"All requesters shall")
    regx13 = re.compile(r"all Requesters shall")
    regx14 = re.compile(r"Requesters shall")
    regx2 = re.compile(r"A Responder")
    regx20 = re.compile(r"Responder shall")
    regx21 = re.compile(r"All Responders shall")
    regx22 = re.compile(r"All responders shall")
    regx23 = re.compile(r"all Responders shall")
    regx24 = re.compile(r"Responders shall")
    regx3 = re.compile(r"A component")
    regx30 = re.compile(r"component shall")
    regx31 = re.compile(r"components shall")
    regx4 = re.compile(r" a Requester")
    regx5 = re.compile(r" a Responder")
    regx6 = re.compile(r" a component")

    #Transparent Router
    regx1000 = re.compile(r"TR shall")
    regx1001 = re.compile(r"(TR) shall")
    regx1002 = re.compile(r"Transparent Router shall")
    regx1003 = re.compile(r"transparent router shall")

    #Accelerators
    regx2000 = re.compile(r"accelerator shall")
    regx2001 = re.compile(r"accelerators shall")

    #Management
    regx7 = re.compile(r"Management shall")
    #General Request Packets
    regx8 = re.compile(r"Request packet shall")
    regx80 = re.compile(r"request packet shall")
    regx81 = re.compile(r"request shall")
    regx82 = re.compile(r"requests shall")
    regx83 = re.compile(r"request packets shall")
    #General Response Packets
    regx9 = re.compile(r"Response packet shall")
    regx90 = re.compile(r"response packet shall")
    regx91 = re.compile(r"response shall")
    regx92 = re.compile(r"responses shall")
    regx93 = re.compile(r"response packets shall")

    #Unclassified packet
    regx70 = re.compile(r"packets shall")
    regx71 = re.compile(r"packet shall")

    #Unclassified Packet or Control Space Field?
    regx60 = re.compile(r"field shall")
    regx61 = re.compile(r"fields shall")

    #Switches
    regx50 = re.compile(r"switch shall")
    regx51 = re.compile(r"switches shall")
    regx52 = re.compile(r"Switches shall")

    #Software
    regx40 = re.compile(r"software shall")
    regx41 = re.compile(r"Software shall")

    #Link / Interface
    regx42 = re.compile(r"link shall")
    regx43 = re.compile(r"links shall")
    regx44 = re.compile(r"Links shall")
    regx45 = re.compile(r"interface shall")
    regx46 = re.compile(r"interfaces shall")

    regx100 = re.compile(r"R-Key")
    regx200 = re.compile(r"LPD")
    regx300 = re.compile(r"LPH")
    regx400 = re.compile(r"field shall")
    if(regx1.search(t)):
        return "REQUESTER"
    if(regx10.search(t)):
        return "REQUESTER"
    if(regx11.search(t)):
        return "REQUESTER"
    if(regx12.search(t)):
        return "REQUESTER"
    if(regx13.search(t)):
        return "REQUESTER"
    if(regx14.search(t)):
        return "REQUESTER"
    elif(regx4.search(t)):
        return "REQUESTER"
    elif(regx2.search(t)):
        return "RESPONDER"
    if(regx20.search(t)):
        return "RESPONDER"
    if(regx21.search(t)):
        return "RESPONDER"
    if(regx22.search(t)):
        return "RESPONDER"
    if(regx23.search(t)):
        return "RESPONDER"
    if(regx24.search(t)):
        return "RESPONDER"
    elif(regx5.search(t)):
        return "RESPONDER"
    elif(regx3.search(t)):
        return "COMPONENT"
    if(regx30.search(t)):
        return "COMPONENT"
    elif(regx31.search(t)):
        return "COMPONENT"
    elif(regx6.search(t)):
        return "COMPONENT"
    elif(regx1000.search(t)):
        return "TRANS_ROUTER"
    elif(regx1001.search(t)):
        return "TRANS_ROUTER"
    elif(regx1002.search(t)):
        return "TRANS_ROUTER"
    elif(regx1003.search(t)):
        return "TRANS_ROUTER"
    elif(regx2000.search(t)):
        return "ACCELERATOR"
    elif(regx2001.search(t)):
        return "ACCELERATOR"
    if(regx7.search(t)):
        return "MANAGEMENT"
    if(regx8.search(t)):
        return "RQ_PACKET"
    elif(regx80.search(t)):
        return "RQ_PACKET"
    elif(regx81.search(t)):
        return "RQ_PACKET"
    elif(regx82.search(t)):
        return "RQ_PACKET"
    elif(regx83.search(t)):
        return "RQ_PACKET"
    elif(regx9.search(t)):
        return "RP_PACKET"
    elif(regx90.search(t)):
        return "RP_PACKET"
    elif(regx91.search(t)):
        return "RP_PACKET"
    elif(regx92.search(t)):
        return "RP_PACKET"
    elif(regx93.search(t)):
        return "RP_PACKET"
    elif(regx70.search(t)):
        return "SOME_PACKET"
    elif(regx71.search(t)):
        return "SOME_PACKET"
    elif(regx60.search(t)):
        return "SOME_FIELD"
    elif(regx61.search(t)):
        return "SOME_FIELD"
    elif(regx50.search(t)):
        return "SWITCH"
    elif(regx51.search(t)):
        return "SWITCH"
    elif(regx52.search(t)):
        return "SWITCH"
    elif(regx40.search(t)):
        return "SOFTWARE"
    elif(regx41.search(t)):
        return "SOFTWARE"
    elif(regx42.search(t)):
        return "LINK_LAYER"
    elif(regx43.search(t)):
        return "LINK_LAYER"
    elif(regx44.search(t)):
        return "LINK_LAYER"
    elif(regx45.search(t)):
        return "INTERFACE"
    elif(regx46.search(t)):
        return "INTERFACE"
    elif(regx100.search(t)):
        return "R_KEY"
    elif(regx200.search(t)):
        return "LPD"
    elif(regx300.search(t)):
        return "LPH"
    else:
        return "UNKNOWN"

#Find Assertions through Paragharph Text
def docx_find(doc_obj, regex, sheet, ix):
    snlp = open('snlp.txt','w')
    n = open('np.txt','w')
    u = open('unknown.txt','w')
    a = open('assertions.txt','w')
    num_h1 = 0
    num_h2 = 0
    num_h3 = 0
    num = 0
    init_count = 0
    id_string = ""
    component = ""

    for p in doc_obj.paragraphs:
        if p.style.name == 'Heading 1':
            init_count = init_count + 1
            if init_count > 4:
                num_h1 = num_h1 + 1
                num_h2 = 0
                num_h3 = 0
                num = 0
                #print("SECTION: "+str(num_h1)+"."+str(num_h2)+"."+str(num_h3)+": "+p.text)
        if p.style.name == 'Heading 2':
            if init_count > 3:
                num_h2 = num_h2 + 2
                num_h3 = 0
                num = 0
                #print("SUB-SECTION"+str(num_h1)+"."+str(num_h2)+"."+str(num_h3)+": "+p.text)
        if p.style.name == 'Heading 3':
            if init_count > 3:
                num_h3 = num_h3 + 1
                num = 0
                #print("SUB-SUB-SECTION"+str(num_h1)+"."+str(num_h2)+"."+str(num_h3)+": "+p.text)

        if regex.search(p.text):
            if init_count > 3:
                blob = TextBlob(p.text)
                for sentence in blob.sentences:
                    if regex.search(str(sentence)):
                        num = num + 1
                        ix = ix + 1
                        component = find_component(str(sentence));
                        id_string = "ID:"+str(num_h1)+"."+str(num_h2)+"."+str(num_h3)+"."+str(num)+":"
                        as_string = id_string+"\t"+component+":\t"+str(sentence)+"\n"
                        print(as_string)
                        #Save to Excel Spreadsheet
                        sheet.write(ix, 0, id_string)
                        sheet.write(ix, 1, str(sentence)) # save assertion sentence
                        sheet.write(ix, 2, p.text) # save entire paragraph for context
                        sheet.write(ix, 3, component) # save entire paragraph for context
                        if component != "UNKNOWN":
                            a.write(as_string)
                        elif component == "UNKNOWN":
                            u.write(as_string)

                #Highlight text in the paragraph that contains the assertion
                for run in p.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    #Close files and connection to NLP
    u.close()
    a.close()
    snlp.close()
    return ix

def docx_tab_find(doc_obj, regex, sheet, ix):
    table_num = 0
    num = 0
    for table in doc_obj.tables:
        table_num = table_num + 1
        num = 0
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if regex.search(p.text):
                        num = num + 1
                        print("ID:T."+str(table_num)+"."+str(num)+":\t"+p.text)
                        id_string = "ID:T"+str(table_num)+"."+str(num)
                        #Save to Excel Spreadsheet
                        ix = ix + 1
                        sheet.write(ix, 0, id_string)
                        sheet.write(ix, 1, p.text)
                        for run in p.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return ix



SOURCE_FILE = "filename.docx"
HIGHLIGHTED_FILE = "filename.docx"
SPREADSHEET_FILE = "checklist.txt.docx"

ix_num = 0
book = xlwt.Workbook()
sh = book.add_sheet("Checklist Items")
col1_name = 'CHECKLIST ID'
col2_name = 'DESCRIPTION'
sh.write(0, 0, col1_name)
sh.write(0, 1, col2_name)
regex1 = re.compile(r"shall")
filename = SOURCE_FILE
doc = Document(filename)
ix_num = docx_find(doc, regex1, sh, ix_num)
ix_num = docx_tab_find(doc, regex1, sh, ix_num)
doc.save(HIGHLIGHTED_FILE)
book.save(SPREADSHEET_FILE)

